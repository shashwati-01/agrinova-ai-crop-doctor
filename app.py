import streamlit as st

st.set_page_config(
    page_title="AgriNova",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<link href="https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;500;700&family=Manrope:wght@400;500;700;800&display=swap" rel="stylesheet">
""", unsafe_allow_html=True)

import io
import numpy as np
import pandas as pd
from docx import Document
from lime.lime_tabular import LimeTabularExplainer
from sklearn.ensemble import GradientBoostingClassifier, RandomForestClassifier


def load_css(file_name):
    with open(file_name) as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)


try:
    load_css("style.css")
except Exception:
    st.warning("style.css not found. Please ensure it is in the same folder.")


@st.cache_resource
def train_models():
    crop = pd.read_csv("crop.csv")
    soil = pd.read_csv("soil.csv")
    X_crop = crop.drop("label", axis=1)
    y_crop = crop["label"]

    crop_model = RandomForestClassifier(random_state=42)
    crop_model.fit(X_crop, y_crop)

    X_soil = soil[["N", "P", "K", "pH"]]
    y_soil = soil["Output"]

    soil_model = GradientBoostingClassifier(random_state=42)
    soil_model.fit(X_soil, y_soil)
    return crop_model, soil_model, X_crop


crop_model, soil_model, X_crop = train_models()

with st.sidebar:
    st.markdown('<div class="sidebar-shell">', unsafe_allow_html=True)
    st.markdown("## Agro Inputs")
    st.markdown(
        """
        <div class="sidebar-intro">
            Configure live soil and climate conditions to run the AI diagnosis.
        </div>
        """,
        unsafe_allow_html=True
    )
    st.info("**Ideal Range**\n\nN: 40–80  \nP: 40–60  \nK: 40–80  \npH: 5.5–7.5")

    N = st.slider("Nitrogen (N)", 0, 140, 50)
    P = st.slider("Phosphorus (P)", 0, 140, 50)
    K = st.slider("Potassium (K)", 0, 200, 50)
    temperature = st.slider("Temperature (°C)", 0, 50, 25)
    humidity = st.slider("Humidity (%)", 0, 100, 50)
    pH = st.slider("pH Level", 0.0, 14.0, 7.0)
    rainfall = st.slider("Rainfall (mm)", 0, 2000, 100)
    st.markdown("</div>", unsafe_allow_html=True)

crop_input = np.array([[N, P, K, temperature, humidity, pH, rainfall]])
soil_input = np.array([[N, P, K, pH]])

st.markdown('<div class="page-shell">', unsafe_allow_html=True)
st.markdown(
    """
    <div class="hero-card">
        <div class="hero-grid">
            <div class="hero-copy">
                <div class="hero-badge">Smart Farming Intelligence • ML Powered • Executive Dashboard</div>
                <div class="hero-title">AgriNova</div>
                <div class="hero-subtitle">
                    Precision agriculture analytics platform for crop recommendation, soil health diagnostics,
                    nutrient planning, and explainable AI decisions.
                </div>
                <div class="hero-pills">
                    <span>Crop Intelligence</span>
                    <span>Soil Risk Scoring</span>
                    <span>Explainable AI</span>
                </div>
            </div>
            <div class="hero-panel">
                <div class="hero-panel-label">Live Input Snapshot</div>
                <div class="hero-panel-grid">
                    <div><span>N</span><strong>{}</strong></div>
                    <div><span>P</span><strong>{}</strong></div>
                    <div><span>K</span><strong>{}</strong></div>
                    <div><span>pH</span><strong>{:.1f}</strong></div>
                </div>
                <div class="hero-panel-foot">
                    Temperature {}°C • Humidity {}% • Rainfall {} mm
                </div>
            </div>
        </div>
    </div>
    """.format(N, P, K, pH, temperature, humidity, rainfall),
    unsafe_allow_html=True
)

st.markdown(
    """
    <div class="intro-strip">
        <div class="intro-item">
            <span class="intro-kicker">Model Stack</span>
            <strong>Random Forest + Gradient Boosting</strong>
        </div>
        <div class="intro-item">
            <span class="intro-kicker">Decision Layer</span>
            <strong>LIME-based Explainability</strong>
        </div>
        <div class="intro-item">
            <span class="intro-kicker">Output</span>
            <strong>Interactive + Word Report Export</strong>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

if st.button("Run Comprehensive Diagnosis"):
    proba = crop_model.predict_proba(crop_input)[0]
    labels = crop_model.classes_
    idx = np.argmax(proba)
    crop_pred = labels[idx]
    confidence = proba[idx] * 100

    col1, col2 = st.columns(2, gap="large")

    with col1:
        st.markdown(
            f"""
            <div class="metric-container metric-primary">
                <div class="metric-label">Top Recommended Crop</div>
                <div class="metric-value">{crop_pred}</div>
                <div class="metric-sub">Best match based on current soil and climate inputs</div>
            </div>
            """,
            unsafe_allow_html=True
        )

    with col2:
        st.markdown(
            f"""
            <div class="metric-container metric-accent">
                <div class="metric-label">Confidence Score</div>
                <div class="metric-value">{confidence:.2f}%</div>
                <div class="metric-sub">Model certainty for the top crop recommendation</div>
            </div>
            """,
            unsafe_allow_html=True
        )

    st.markdown('<div class="spacer-24"></div>', unsafe_allow_html=True)

    low_col1, low_col2 = st.columns(2, gap="large")

    with low_col1:
        st.markdown(
            """
            <div class="section-card">
                <div class="section-header">
                    <div>
                        <div class="section-overline">Diagnostics</div>
                        <div class="section-title">Soil Status & Warnings</div>
                        <div class="section-subtitle">Live evaluation of soil quality and critical alerts</div>
                    </div>
                </div>
            """,
            unsafe_allow_html=True
        )

        soil_pred_num = soil_model.predict(soil_input)[0]
        soil_map = {0: "Poor", 1: "Medium", 2: "Good"}
        soil_pred = soil_map.get(soil_pred_num, "Unknown")

        if soil_pred == "Poor":
            st.error(f"Soil Condition: {soil_pred}")
        elif soil_pred == "Medium":
            st.warning(f"Soil Condition: {soil_pred}")
        else:
            st.success(f"Soil Condition: {soil_pred}")

        if rainfall < 20:
            st.warning("Low rainfall detected")

        if pH < 5 or pH > 8:
            st.warning("pH out of optimal range")

        st.markdown("</div>", unsafe_allow_html=True)

    with low_col2:
        st.markdown(
            """
            <div class="section-card">
                <div class="section-header">
                    <div>
                        <div class="section-overline">Guidance</div>
                        <div class="section-title">Fertilizer Recommendation</div>
                        <div class="section-subtitle">Nutrient actions to improve crop suitability</div>
                    </div>
                </div>
            """,
            unsafe_allow_html=True
        )

        fert_list = []
        if N < 40:
            fert_list.append("Nitrogen fertilizer (Urea)")
        if P < 40:
            fert_list.append("Phosphorus fertilizer (DAP)")
        if K < 40:
            fert_list.append("Potassium fertilizer (MOP)")

        if not fert_list:
            st.success("Soil nutrients are perfectly balanced!")
        else:
            for fert in fert_list:
                st.markdown(
                    f"""
                    <div class="insight-item">Add <strong>{fert}</strong></div>
                    """,
                    unsafe_allow_html=True
                )

        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="spacer-24"></div>', unsafe_allow_html=True)

    vis_col1, vis_col2 = st.columns(2, gap="large")

    with vis_col1:
        st.markdown(
            """
            <div class="section-card chart-card">
                <div class="section-header">
                    <div>
                        <div class="section-overline">Prediction Spread</div>
                        <div class="section-title">Top 3 Suggestions</div>
                        <div class="section-subtitle">Most probable crop recommendations by the model</div>
                    </div>
                </div>
            """,
            unsafe_allow_html=True
        )

        top3 = np.argsort(proba)[-3:][::-1]
        crops = [labels[i] for i in top3]
        probs = [proba[i] * 100 for i in top3]

        df_chart = pd.DataFrame({
            "Crop": crops,
            "Probability (%)": probs
        })

        st.bar_chart(df_chart.set_index("Crop"))
        st.markdown("</div>", unsafe_allow_html=True)

    with vis_col2:
        st.markdown(
            """
            <div class="section-card">
                <div class="section-header">
                    <div>
                        <div class="section-overline">Explainability</div>
                        <div class="section-title">AI Decision Insight (LIME)</div>
                        <div class="section-subtitle">Top feature contributions influencing the recommendation</div>
                    </div>
                </div>
            """,
            unsafe_allow_html=True
        )

        explainer = LimeTabularExplainer(
            X_crop.values,
            feature_names=X_crop.columns,
            class_names=labels,
            mode="classification"
        )

        exp = explainer.explain_instance(
            crop_input[0],
            crop_model.predict_proba,
            num_features=5
        )

        for feature, weight in exp.as_list():
            st.markdown(
                f"""
                <div class="insight-item insight-strong"><strong>{feature}</strong> <span>{weight:.3f}</span></div>
                """,
                unsafe_allow_html=True
            )

        st.markdown("</div>", unsafe_allow_html=True)

st.markdown(
    """
    <div class="section-card report-card">
        <div class="report-orb"></div>
        <div class="section-overline">Export</div>
        <div class="section-title">Premium Report</div>
        <div class="section-subtitle">Export your current diagnosis as a clean Word document</div>
    """,
    unsafe_allow_html=True
)

if "crop_pred" in locals():
    doc = Document()

    doc.add_heading("AgriNova Crop Intelligence Report", 0)

    doc.add_heading("Input Parameters", 1)
    doc.add_paragraph(f"Nitrogen (N): {N}")
    doc.add_paragraph(f"Phosphorus (P): {P}")
    doc.add_paragraph(f"Potassium (K): {K}")
    doc.add_paragraph(f"Temperature: {temperature} °C")
    doc.add_paragraph(f"Humidity: {humidity} %")
    doc.add_paragraph(f"pH: {pH}")
    doc.add_paragraph(f"Rainfall: {rainfall} mm")

    doc.add_heading("Crop Recommendation", 1)
    doc.add_paragraph(f"Recommended Crop: {crop_pred}")
    doc.add_paragraph(f"Confidence Score: {confidence:.2f}%")

    doc.add_heading("Soil Health Status", 1)
    doc.add_paragraph(f"Soil Condition: {soil_pred}")

    doc.add_heading("Fertilizer Recommendation", 1)
    if not fert_list:
        doc.add_paragraph("Soil nutrients are balanced.")
    else:
        for fert in fert_list:
            doc.add_paragraph(f"Add {fert}")

    doc.add_heading("Top 3 Crop Suggestions", 1)
    for crop_name, prob in zip(crops, probs):
        doc.add_paragraph(f"{crop_name} - {prob:.2f}%")

    doc.add_heading("AI Decision Insights (LIME)", 1)
    for feature, weight in exp.as_list():
        doc.add_paragraph(f"{feature} -> {weight:.3f}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        "Download Premium Report",
        data=buffer,
        file_name="AI_Crop_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.markdown("</div>", unsafe_allow_html=True)
st.markdown("</div>", unsafe_allow_html=True)